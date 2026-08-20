"""Genera el documento Word final de resultados tecnicos de la tesis."""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import os
import zipfile
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


DEFAULT_ROOT = Path(os.environ.get("THESIS_PROJECT_ROOT", Path(__file__).resolve().parents[2]))
NAVY = "17365D"
BLUE = "2F5597"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F4F7"
MID_GRAY = "667085"
WHITE = "FFFFFF"


def read_csv(path: Path) -> list[dict[str, str]]:
    if not path.exists():
        return []
    with path.open("r", encoding="utf-8-sig", newline="") as stream:
        return list(csv.DictReader(stream))


def fnum(value: object, digits: int = 4, default: str = "N/D") -> str:
    try:
        return f"{float(value):.{digits}f}"
    except (TypeError, ValueError):
        return default


def inum(value: object, default: str = "N/D") -> str:
    try:
        return str(int(float(value)))
    except (TypeError, ValueError):
        return default


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for block in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tc_pr.append(node)
    node.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=80, bottom=70, end=80) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = tc_mar.find(qn(f"w:{margin}"))
        if element is None:
            element = OxmlElement(f"w:{margin}")
            tc_mar.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_indent(table, twips: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    indent = tbl_pr.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:w"), str(twips))
    indent.set(qn("w:type"), "dxa")


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Página ")
    run.font.size = Pt(9)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr_text, fld_char2])


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    heading_specs = {
        "Title": (28, NAVY, 0, 12),
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (11.5, NAVY, 9, 4),
    }
    for name, (size, color, before, after) in heading_specs.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    if "Figure Caption" not in [style.name for style in doc.styles]:
        style = doc.styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Caption"]
    figure_caption = doc.styles["Figure Caption"]
    figure_caption.font.name = "Calibri"
    figure_caption.font.size = Pt(9)
    figure_caption.font.italic = True
    figure_caption.font.color.rgb = RGBColor.from_string(MID_GRAY)
    figure_caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure_caption.paragraph_format.space_after = Pt(8)

    header = section.header.paragraphs[0]
    header.text = "RESULTADOS TÉCNICOS FINALES | SISTEMA DE VISIÓN ARTIFICIAL HEPÁTICA"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.color.rgb = RGBColor.from_string(MID_GRAY)
    add_page_number(section.footer.paragraphs[0])


def add_cover(doc: Document, timestamp: str) -> None:
    doc.add_paragraph("UNIVERSIDAD INTERNACIONAL DEL ECUADOR", style="Subtitle").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\n")
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Development of an Artificial Vision System for the Standardized Acquisition of Clinically Relevant Hepatic Ultrasound Images")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("RESULTADOS TÉCNICOS FINALES DEL SISTEMA DE VISIÓN ARTIFICIAL HEPÁTICA")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor.from_string(BLUE)
    doc.add_paragraph("\n\n")
    for text in (
        "Autor: Julián Guapaz Pozo",
        "Institución: Universidad Internacional del Ecuador",
        f"Fecha de congelamiento técnico: {datetime.strptime(timestamp, '%Y%m%d_%H%M').strftime('%d/%m/%Y %H:%M')}",
        "Documento de consolidación de evidencia reproducible",
    ):
        paragraph = doc.add_paragraph(text)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.paragraph_format.space_before = Pt(28)
    note.add_run(
        "Prototipo experimental para procesamiento secuencial de videos ecográficos previamente grabados. "
        "No realiza diagnóstico ni constituye validación clínica."
    ).italic = True
    doc.add_page_break()


def add_text(doc: Document, text: str, bold_lead: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_lead and text.startswith(bold_lead):
        paragraph.add_run(bold_lead).bold = True
        paragraph.add_run(text[len(bold_lead):])
    else:
        paragraph.add_run(text)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.add_run(item)


def add_note(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    shade(cell, LIGHT_BLUE)
    set_cell_margins(cell, 100, 130, 100, 130)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run("Nota metodológica. ")
    run.bold = True
    paragraph.add_run(text)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_table(doc: Document, caption: str, headers: list[str], rows: list[list[object]], widths: list[float] | None = None) -> None:
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.paragraph_format.keep_with_next = True
    run = caption_paragraph.add_run(caption)
    run.bold = True
    run.font.size = Pt(9.5)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_indent(table)
    header_row = table.rows[0]
    set_repeat_header(header_row)
    for index, header in enumerate(headers):
        cell = header_row.cells[index]
        cell.text = str(header)
        shade(cell, LIGHT_GRAY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8.5)
        set_cell_margins(cell)
    for row_data in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_data):
            cells[index].text = str(value)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(8.2)
            set_cell_margins(cells[index])
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_figure(doc: Document, path: Path, caption: str, width: float = 6.35) -> bool:
    if not path.exists():
        add_note(doc, f"La figura prevista no estaba disponible en el paquete al generar el documento: {path.name}.")
        return False
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width))
    caption_paragraph = doc.add_paragraph(caption, style="Figure Caption")
    caption_paragraph.paragraph_format.keep_with_next = False
    return True


def find_row(rows: list[dict[str, str]], **criteria: str) -> dict[str, str]:
    for row in rows:
        if all(row.get(key, "").strip().lower() == value.lower() for key, value in criteria.items()):
            return row
    return {}


def build_document(root: Path, timestamp: str) -> tuple[Path, dict[str, object]]:
    evidence = root / "outputs" / f"TESIS_EVIDENCIA_FINAL_{timestamp}"
    figures = evidence / "07_figuras_para_tesis"
    output_dir = evidence / "10_documento_word_final"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / "RESULTADOS_TECNICOS_TESIS_FINAL.docx"

    exp = root / "outputs" / "experimental_segmentation_pipeline"
    dino = root / "outputs" / "dino_experimental" / "binary_improvement"
    models = exp / "final_models"
    metrics = read_csv(exp / "reports" / "all_metrics.csv")
    benchmarks = read_csv(exp / "reports" / "all_benchmarks.csv")
    manifest = read_csv(models / "selected_models_manifest.csv")
    p005 = read_csv(exp / "p005_longitudinal_final" / "summary_by_quality.csv")
    dino_internal = read_csv(dino / "reports" / "08_binary_winners_by_view.csv")
    dino_p005 = read_csv(dino / "reports" / "09_binary_p005_metrics.csv")
    dino_benchmark = read_csv(dino / "reports" / "12_binary_dino_inference_benchmark.csv")
    final_benchmark = read_csv(exp / "reports" / "07_inference_benchmark.csv")

    doc = Document()
    configure_document(doc)
    add_cover(doc, timestamp)

    doc.add_heading("1. Propósito del documento", level=1)
    add_text(doc, "Este documento consolida la evidencia técnica final del prototipo de visión artificial desarrollado para apoyar la adquisición estandarizada de ecografías hepáticas. Reúne los modelos congelados, resultados controlados, evaluaciones operacionales, decisiones metodológicas, figuras y limitaciones que sustentan la reproducibilidad del sistema.")

    doc.add_heading("2. Alcance del sistema", level=1)
    add_text(doc, "El producto es un prototipo experimental de asistencia. Procesa videos ecográficos previamente grabados de manera secuencial, frame por frame, para simular el flujo operativo de una adquisición asistida. No captura directamente desde una cámara ni desde un ecógrafo en tiempo real.")
    add_bullets(doc, [
        "No realiza diagnóstico médico, detección de enfermedades ni reemplaza la evaluación profesional.",
        "Genera mensajes de guía, decisiones temporales, capturas confirmadas, mejores candidatos y archivos reproducibles.",
        "La evidencia P005 es funcional u operacional; no constituye validación clínica cuantitativa externa.",
        "Los pesos, umbrales y reglas de decisión se conservaron sin reentrenamiento ni recalibración durante este cierre.",
    ])

    doc.add_heading("3. Flujo general del sistema", level=1)
    add_figure(doc, figures / "fig_01_sistema_completo.png", "Figura 1. Flujo general del prototipo para procesar videos ecográficos previamente grabados.")
    add_text(doc, "La entrada es un video previamente grabado y una vista seleccionada por el operador. La vista longitudinal se analiza mediante tres modelos de segmentación y reglas de calidad; las vistas transversal, oblicua y hepatorrenal emplean DINOv2 y clasificación temporal. La interfaz presenta retroalimentación y conserva CSV, capturas y el mejor frame confirmado o candidato.")

    doc.add_heading("4. Dataset y organización", level=1)
    add_table(doc, "Tabla 1. Organización de los datos empleados.", ["Componente", "Uso", "Observación metodológica"], [
        ["P001-P003", "Desarrollo, división agrupada y evaluación interna", "Frames de videos; riesgo de similitud temporal mitigado con split group_video."],
        ["P004", "Excluido", "No se incorporó en los resultados finales."],
        ["P005", "Prueba externa operacional", "Sin máscaras ground truth; no permite Dice/IoU externo."],
        ["Vistas", "Longitudinal, transversal, oblicua y hepatorrenal", "Calidades clear, medium y blurry son etiquetas nominales por carpeta/video."],
    ], [1.0, 2.0, 3.4])
    add_note(doc, "P005 se interpreta únicamente como evidencia del comportamiento operacional del prototipo ante videos no usados para entrenar. Sin anotaciones externas no es válido presentarlo como validación cuantitativa clínica.")

    doc.add_heading("5. Segmentación longitudinal", level=1)
    doc.add_heading("5.1 Anotación y clases", level=2)
    add_table(doc, "Tabla 2. Clases binarias del pipeline longitudinal.", ["Clase", "Función", "Reto principal"], [
        ["ROI", "Delimita el campo ecográfico útil", "Mantener cobertura estable del abanico ecográfico."],
        ["Hígado", "Segmenta el parénquima hepático", "Variación anatómica, textura y límites difusos."],
        ["LA", "Segmenta la referencia anatómica/lumen", "Estructura pequeña, pocas imágenes positivas y muchas máscaras vacías."],
    ], [0.8, 2.5, 3.1])

    doc.add_heading("5.2 Arquitecturas comparadas", level=2)
    add_table(doc, "Tabla 3. Arquitecturas longitudinales comparadas.", ["Arquitectura", "Encoder/backbone", "Parámetros aprox.", "Justificación"], [
        ["U-Net", "ResNet-34", "24.44 M", "Baseline clásico de segmentación biomédica con conexiones de salto."],
        ["DeepLabV3+", "ResNet-34", "22.44 M", "Contexto multiescala y decodificación de bordes mediante convoluciones atrous."],
        ["SegFormer", "MiT-B0", "3.71 M", "Arquitectura transformer compacta para comparar eficiencia y representación global."],
    ], [1.1, 1.2, 1.1, 3.0])

    doc.add_heading("5.3 Configuración experimental", level=2)
    add_table(doc, "Tabla 4. Hiperparámetros controlados principales.", ["Parámetro", "Valor final/controlado"], [
        ["División", "group_video: train 606, valid 202, test 101"],
        ["Resolución base", "512 × 512, full resize"],
        ["Optimizador", "AdamW"],
        ["Learning rate / weight decay", "0.001 / 0.0001"],
        ["Batch size", "2"],
        ["Épocas máximas", "50"],
        ["Early stopping", "10 épocas para ROI/Hígado; 15 para LA"],
        ["Seed", "42"],
        ["Checkpoint ROI/Hígado", "Dice de validación"],
        ["Checkpoint LA", "combined_la_score = positive_dice − empty_false_positive_rate"],
        ["Muestreo LA seleccionado", "balanced_la, 50 % de positivos efectivos por época"],
    ], [2.2, 4.2])

    doc.add_heading("5.4 Métricas", level=2)
    add_table(doc, "Tabla 5. Métricas de segmentación.", ["Métrica", "Definición"], [
        ["Dice", "2TP / (2TP + FP + FN)"],
        ["IoU", "TP / (TP + FP + FN)"],
        ["Precision", "TP / (TP + FP)"],
        ["Recall", "TP / (TP + FN)"],
        ["F1", "2 · Precision · Recall / (Precision + Recall)"],
        ["combined_la_score", "Dice positivo − tasa de falsos positivos en imágenes vacías"],
    ], [2.0, 4.4])
    add_text(doc, "Para LA, el Dice global puede estar inflado por imágenes vacías correctamente predichas. Por ello, la selección prioriza el Dice en imágenes positivas y penaliza explícitamente las predicciones positivas en ground truth vacío.")

    doc.add_heading("5.5 Resultados controlados", level=2)
    selected_rows = []
    selected_specs = (("ROI", "ROI"), ("Higado", "Hígado"), ("LA", "LA"))
    for manifest_class, display_class in selected_specs:
        selected = next(
            (row for row in manifest if row.get("class_name", "").lower() == manifest_class.lower()), {}
        )
        experiment_name = selected.get("experiment_name", "")
        row = next(
            (
                candidate for candidate in metrics
                if candidate.get("experiment_name", "") == experiment_name and candidate.get("test_dice", "")
            ),
            {},
        )
        selected_rows.append([
            display_class,
            "DeepLabV3+" if selected.get("architecture", "").lower() == "deeplabv3" else "U-Net",
            fnum(row.get("test_dice")),
            fnum(row.get("test_iou")),
            fnum(row.get("test_positive_dice")),
            fnum(row.get("test_empty_gt_false_positive_rate")),
        ])
    add_table(doc, "Tabla 6. Resultados de los modelos longitudinales seleccionados en el test controlado.", ["Clase", "Modelo", "Dice", "IoU", "Dice positivo", "FP vacías"], selected_rows, [0.7, 1.1, 0.8, 0.8, 1.1, 1.0])
    add_figure(doc, figures / "fig_04_comparacion_segmentacion.png", "Figura 2. Comparación de arquitecturas para ROI, Hígado y LA.")

    doc.add_heading("5.6 Ablations", level=2)
    add_text(doc, "Se evaluaron modificaciones de augmentation, resolución y transferencia de aprendizaje. Los ensayos se conservaron aun cuando produjeron resultados negativos, porque permiten justificar la configuración final y mostrar que la selección no se basó únicamente en resultados favorables.")
    add_figure(doc, figures / "fig_05_ablation_resolucion.png", "Figura 3. Ablation de resolución y estrategia de redimensionamiento para LA.")
    add_figure(doc, figures / "fig_06_transfer_learning.png", "Figura 4. Comparación con y sin transferencia de aprendizaje.")
    add_table(doc, "Tabla 7. Hallazgos principales de las ablaciones.", ["Ensayo", "Resultado técnico"], [
        ["Augmentation x4 general", "No mejoró de forma consistente LA; aumentó el costo de entrenamiento y mantuvo inestabilidad."],
        ["positive_x4", "Aumentó positivos efectivos, pero obtuvo Dice test 0.4355; no fue seleccionado."],
        ["ROI crop 128", "Aumentó velocidad (~110.6 FPS individual), pero redujo el desempeño de LA (Dice 0.4709)."],
        ["Padding 512", "Colapsó con frecuencia a máscara vacía; Dice test 0.4554."],
        ["Pretraining ImageNet LA", "Mejoró el resultado controlado de U-Net LA hasta Dice 0.8425 y fue seleccionado."],
    ], [1.8, 4.6])

    doc.add_heading("5.7 Modelos longitudinales seleccionados", level=2)
    hash_by_class = {}
    for filename, class_name in (("best_roi_model.pth", "ROI"), ("best_higado_model.pth", "Hígado"), ("best_la_model.pth", "LA")):
        path = models / filename
        hash_by_class[class_name] = sha256(path) if path.exists() else "N/D"
    add_table(doc, "Tabla 8. Modelos longitudinales finales congelados.", ["Clase", "Arquitectura", "Pretraining/muestreo", "SHA-256"], [
        ["ROI", "DeepLabV3+ ResNet-34", "Sin pretraining", hash_by_class["ROI"]],
        ["Hígado", "DeepLabV3+ ResNet-34", "ImageNet, fine-tuning completo", hash_by_class["Hígado"]],
        ["LA", "U-Net ResNet-34", "ImageNet + balanced_la", hash_by_class["LA"]],
    ], [0.7, 1.6, 1.8, 2.3])

    doc.add_heading("5.8 Regla longitudinal final", level=2)
    add_figure(doc, figures / "fig_02_pipeline_longitudinal.png", "Figura 5. Pipeline longitudinal congelado y compartido por inferencia y GUI.")
    add_bullets(doc, [
        "La región GLCM se define como ROI AND LA dilatada con kernel 15 × 15, coherente con los umbrales offline originales.",
        "El área y la desviación estándar de LA se calculan sobre la máscara postprocesada; el umbral de área de referencia es 874 px a 1024 × 768 y se escala por área del frame.",
        "La regla base exige LA presente, área suficiente, desviación estándar ≤ 43.330219 y entropía GLCM ≤ 6.125722.",
        "La evidencia de borde usa un anillo 9 × 9 y combina brillo, gradiente y contraste local; el modo congelado es regla_lumen_or_border.",
        "La decisión temporal utiliza ventana 5, tres capturas crudas consecutivas para confirmar y cooldown de 10 evaluaciones.",
    ])

    doc.add_heading("5.9 Evaluación operacional P005 longitudinal", level=2)
    p005_rows = [[
        row.get("quality", ""), inum(row.get("frames_processed")), inum(row.get("roi_present_count")),
        inum(row.get("liver_present_count")), inum(row.get("la_present_count")),
        inum(row.get("raw_capture_count")), inum(row.get("stable_capture_count")),
        fnum(row.get("fps_effective"), 2),
    ] for row in p005]
    add_table(doc, "Tabla 9. Resultado operacional P005 longitudinal sin ground truth.", ["Calidad", "Frames", "ROI", "Hígado", "LA", "Captura cruda", "Confirmada", "FPS"], p005_rows, [0.8, 0.7, 0.6, 0.7, 0.6, 1.0, 0.9, 0.7])
    add_figure(doc, figures / "fig_08_p005_longitudinal.png", "Figura 6. Ejemplos P005 longitudinales: ROI e Hígado presentes, referencia LA insuficiente y sin captura confirmada.")
    add_text(doc, "Se evaluaron 102 frames (34 clear, 34 medium y 34 blurry). ROI e Hígado aparecieron en todos; LA apareció en tres frames medium. Ningún candidato alcanzó el umbral de área escalado, por lo que no hubo capturas crudas ni confirmadas. Este resultado evidencia una política conservadora y confirma que LA es el cuello de botella operacional.")
    add_note(doc, "P005 no dispone de máscaras ground truth. No se calcularon Dice ni IoU externos y el resultado no debe presentarse como validación clínica cuantitativa.")

    doc.add_heading("6. DINOv2 para transversal, oblicua y hepatorrenal", level=1)
    doc.add_heading("6.1 Justificación y reformulación", level=2)
    add_text(doc, "Para las vistas no longitudinales se utilizó DINOv2 como extractor auto-supervisado de características visuales. La formulación ternaria inicial no fue suficientemente estable; se adoptó clasificación binaria clear/blurry y se trató medium como zona de incertidumbre, con umbrales de abstención.")
    add_figure(doc, figures / "fig_03_pipeline_dino.png", "Figura 7. Pipeline DINOv2 temporal para las vistas transversal, oblicua y hepatorrenal.")
    doc.add_heading("6.2 Backbones, clasificadores y temporalidad", level=2)
    add_table(doc, "Tabla 10. Bundles DINOv2 finales usados por la GUI.", ["Vista", "Preprocesamiento", "Ventana", "Clasificador"], [
        ["Transversal", "small + fan_crop", "5 embeddings", "Random Forest"],
        ["Oblicua", "small + fan_crop", "5 embeddings", "Regresión logística"],
        ["Hepatorrenal", "small + full", "5 embeddings", "k-NN"],
    ], [1.1, 1.8, 1.2, 2.2])
    add_text(doc, "Los umbrales operativos congelados fueron adjust < 0.35, doubtful entre 0.35 y 0.65 y capture ≥ 0.65. La captura requiere confirmación temporal y cooldown para evitar guardar secuencias redundantes.")
    doc.add_heading("6.3 Resultados internos y P005", level=2)
    internal_rows = [[
        row.get("view", ""), row.get("embedding_variant", ""), row.get("classifier", ""),
        fnum(row.get("mean_video_f1_macro")), fnum(row.get("mean_video_balanced_accuracy")),
    ] for row in dino_internal]
    if internal_rows:
        add_table(doc, "Tabla 11. Resumen LOPO interno de la configuración binaria seleccionada.", ["Vista", "Variante", "Clasificador", "F1 macro", "Balanced acc."], internal_rows, [1.0, 1.3, 1.5, 0.9, 1.1])
    doc.add_page_break()
    p005_dino_rows = [[
        row.get("view", ""), fnum(row.get("video_f1_macro")),
        fnum(row.get("action_capture_rate_clear")), fnum(row.get("action_false_capture_rate_blurry")),
        fnum(row.get("action_anchor_abstention_rate")),
    ] for row in dino_p005]
    if p005_dino_rows:
        add_table(doc, "Tabla 12. Resumen operacional P005 de las vistas DINOv2.", ["Vista", "F1 video", "Capture clear", "False capture blurry", "Abstención"], p005_dino_rows, [1.1, 0.9, 1.1, 1.4, 1.0])
    add_figure(doc, figures / "fig_09_dino_vistas.png", "Figura 8. Distribución de acciones DINOv2 en P005 por vista.")
    add_note(doc, "Las probabilidades DINOv2 no fueron calibradas formalmente y P005 no constituye una validación clínica. Los resultados LOPO internos se interpretan en el contexto de pocos pacientes y videos correlacionados.")

    doc.add_heading("7. GUI unificada", level=1)
    add_figure(doc, figures / "fig_10_gui_final.png", "Figura 9. Evidencia de la GUI Streamlit procesando videos previamente grabados y guardando resultados de sesión.")
    add_bullets(doc, [
        "GUI desarrollada con Streamlit para cargar MP4, AVI, MOV o MKV, o indicar una ruta local.",
        "Procesa videos previamente grabados; no captura directamente desde cámara ni ecógrafo en tiempo real.",
        "Crea una carpeta independiente por sesión y admite longitudinal, transversal, oblicua y hepatorrenal.",
        "Longitudinal usa segmentación ROI/Hígado/LA; las otras vistas usan DINOv2 y clasificador binario temporal.",
        "Usa caché de modelos y permite configurar stride, máximo de frames, frecuencia de actualización, cooldown y CPU/GPU.",
        "Muestra frame original y overlay o campo ecográfico limpio, junto con mensaje y decisión.",
        "Guarda frame_results.csv, session_summary, capturas confirmadas y mejor frame confirmado o candidato.",
        "El navegador DINO permite revisar cualquiera de los frames analizados; no selecciona siempre el último frame.",
        "Es una interfaz experimental de asistencia y no una herramienta diagnóstica.",
    ])

    doc.add_heading("8. Benchmark de inferencia", level=1)
    component_labels = {
        "ROI": "ROI DeepLabV3+", "Higado": "Hígado DeepLabV3+", "LA": "LA U-Net",
        "pipeline_3_models_total": "Pipeline longitudinal completo",
    }
    bench_rows = []
    for row in final_benchmark:
        component = row.get("component", "")
        if component in component_labels:
            bench_rows.append([
                component_labels[component], fnum(row.get("mean_ms_per_frame"), 2),
                fnum(row.get("fps"), 2), "Sí" if row.get("meets_30_fps", "").lower() == "true" else "No",
            ])
    for row in dino_benchmark:
        view = row.get("view", "").capitalize()
        bench_rows.append([
            f"DINO {view}, cada frame", fnum(row.get("mean_ms_per_evaluated_frame"), 2),
            fnum(row.get("fps_if_every_frame"), 2), "No",
        ])
    add_table(doc, "Tabla 13. Benchmark de inferencia en el equipo experimental.", ["Modelo/pipeline", "ms/frame", "FPS", "≥30 FPS"], bench_rows, [3.2, 1.0, 0.8, 0.9])
    add_figure(doc, figures / "fig_07_benchmark_fps.png", "Figura 10. FPS de los modelos longitudinales individuales y del pipeline de tres redes; referencia de 30 FPS.")
    add_text(doc, "Cada red longitudinal individual supera 30 FPS en el benchmark aislado, pero las tres redes en serie alcanzan aproximadamente 20.29 FPS y no cumplen 30 FPS sin reducción de frecuencia. La GUI utiliza stride para sostener un flujo interactivo sobre videos grabados.")

    doc.add_heading("9. Resultados negativos y decisiones metodológicas", level=1)
    add_bullets(doc, [
        "Los resultados preliminares de cinco épocas se conservaron como antecedentes, pero no se usaron como evidencia final robusta.",
        "El recorte ROI a 128 × 128 fue rápido, pero deterioró LA y elevó el riesgo de falsos positivos/vacíos.",
        "La estrategia original_or_padding mostró colapso hacia máscara vacía.",
        "La augmentation x4, incluida la variante solo positiva, no mejoró LA de manera defendible.",
        "La clasificación DINO ternaria fue inestable y motivó la reformulación binaria con abstención.",
        "P005 longitudinal no produjo capturas confirmadas porque los candidatos LA no alcanzaron el área mínima.",
    ])

    doc.add_heading("10. Limitaciones", level=1)
    add_bullets(doc, [
        "Número reducido de pacientes y correlación temporal entre frames de video.",
        "P005 sin ground truth, por lo que solo permite evaluación operacional.",
        "Etiquetas clear, medium y blurry nominales por carpeta/video.",
        "Test interno limitado y ausencia de estudio formal con usuarios no expertos.",
        "No se realizó validación clínica ni se afirma generalización clínica.",
        "Umbrales principalmente derivados del conjunto de desarrollo.",
        "El pipeline longitudinal completo no alcanza 30 FPS sin stride.",
        "Probabilidades DINOv2 no calibradas formalmente.",
        "La GUI no captura directamente desde ecógrafo o cámara en vivo.",
    ])

    doc.add_heading("11. Conclusión técnica", level=1)
    add_text(doc, "Se desarrolló y congeló un prototipo funcional y reproducible que integra segmentación longitudinal, DINOv2 para otras vistas y una GUI Streamlit. El sistema procesa videos previamente grabados, presenta retroalimentación y conserva evidencia estructurada por sesión. LA continúa siendo el cuello de botella longitudinal. El prototipo queda preparado para una evaluación futura con más pacientes, anotaciones ground truth, estudio con usuarios y, si se desea, integración posterior con un flujo real de video.")

    doc.add_heading("12. Archivos entregables", level=1)
    backup = root / f"TESIS_VERSION_FINAL_{timestamp}"
    add_table(doc, "Tabla 14. Entregables del cierre técnico.", ["Entregable", "Ubicación relativa"], [
        ["Backup congelado", backup.name],
        ["Paquete de evidencia", f"outputs/{evidence.name}"],
        ["Modelos y hashes", "02_modelos_y_hashes/"],
        ["Reportes y CSV", "01_metodologia/ a 06_resultados_p005/"],
        ["Figuras", "07_figuras_para_tesis/"],
        ["Overlays y casos", "08_overlays_y_casos/"],
        ["Anexos reproducibilidad", "09_anexos_reproducibilidad/"],
        ["Documento Word", "10_documento_word_final/RESULTADOS_TECNICOS_TESIS_FINAL.docx"],
    ], [2.1, 4.3])

    doc.add_heading("13. Referencias base", level=1)
    references = [
        "Ronneberger, O., Fischer, P., & Brox, T. (2015). U-Net: Convolutional Networks for Biomedical Image Segmentation.",
        "Chen, L.-C., Zhu, Y., Papandreou, G., Schroff, F., & Adam, H. (2018). Encoder-Decoder with Atrous Separable Convolution for Semantic Image Segmentation.",
        "Xie, E., Wang, W., Yu, Z., Anandkumar, A., Alvarez, J. M., & Luo, P. (2021). SegFormer: Simple and Efficient Design for Semantic Segmentation with Transformers.",
        "Oquab, M., et al. (2023). DINOv2: Learning Robust Visual Features without Supervision.",
        "He, K., Zhang, X., Ren, S., & Sun, J. (2016). Deep Residual Learning for Image Recognition.",
        "Haralick, R. M., Shanmugam, K., & Dinstein, I. (1973). Textural Features for Image Classification.",
        "Lin, T.-Y., et al. (2014). Microsoft COCO: Common Objects in Context.",
        "Loshchilov, I., & Hutter, F. (2019). Decoupled Weight Decay Regularization.",
    ]
    for reference in references:
        paragraph = doc.add_paragraph(reference)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.paragraph_format.left_indent = Inches(0.25)
        paragraph.paragraph_format.space_after = Pt(4)

    doc.save(output_path)

    with zipfile.ZipFile(output_path) as archive:
        names = archive.namelist()
        embedded_images = len([name for name in names if name.startswith("word/media/")])
    validation = {
        "generated_at": datetime.now().isoformat(timespec="seconds"),
        "path": str(output_path),
        "size_bytes": output_path.stat().st_size,
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "inline_shapes": len(doc.inline_shapes),
        "embedded_images": embedded_images,
        "critical_checks": {
            "nonempty": output_path.stat().st_size > 100_000,
            "has_tables": len(doc.tables) > 0,
            "has_images": embedded_images > 0,
            "zip_valid": True,
        },
    }
    (output_dir / "word_generation_validation.json").write_text(
        json.dumps(validation, ensure_ascii=False, indent=2), encoding="utf-8"
    )
    return output_path, validation


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--project-root", type=Path, default=DEFAULT_ROOT)
    parser.add_argument("--timestamp", required=True, help="Timestamp YYYYMMDD_HHMM del paquete congelado")
    args = parser.parse_args()
    output_path, validation = build_document(args.project_root.resolve(), args.timestamp)
    print(f"Word: {output_path}")
    print(json.dumps(validation, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
