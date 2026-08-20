"""Genera un documento Word tecnico con los resultados longitudinales."""

from __future__ import annotations

import os
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


TESIS_ROOT = Path(os.environ.get("THESIS_PROJECT_ROOT", Path(__file__).resolve().parents[1]))
OUTPUTS_ROOT = TESIS_ROOT / "outputs"
REPORTS_ROOT = OUTPUTS_ROOT / "reports"
FIGURES_ROOT = OUTPUTS_ROOT / "figures"
DOCX_PATH = REPORTS_ROOT / "reporte_tecnico_pipeline_longitudinal_actualizado_v2.docx"

ANALYSIS_CSV = REPORTS_ROOT / "longitudinal_acceptability_analysis.csv"
EXAMPLES_CSV = REPORTS_ROOT / "longitudinal_acceptability_examples.csv"
VISUAL_CASES_CSV = REPORTS_ROOT / "longitudinal_visual_case_examples.csv"
THRESHOLD_CSV = REPORTS_ROOT / "threshold_summary.csv"
QUALITY_GROUP_CSV = REPORTS_ROOT / "quality_group_summary.csv"
ROBOFLOW_AUDIT_CSV = REPORTS_ROOT / "roboflow_dataset_audit.csv"
MASK_REPORT_CSV = REPORTS_ROOT / "mask_generation_report.csv"
METRICS_CSV = OUTPUTS_ROOT / "metrics" / "glcm_longitudinal_metrics_with_thresholds.csv"
VISUAL_CASES_FIGURE = FIGURES_ROOT / "casos_representativos_longitudinal.png"


def set_cell_shading(cell, fill: str) -> None:
    """Aplica color de fondo a una celda."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    """Escribe texto con formato consistente en una celda."""
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(8.5)
    paragraph.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_geometry(table, widths_in: list[float]) -> None:
    """Fija anchos de tabla/celdas para evitar autofit inestable."""
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths_in):
                cell.width = Inches(widths_in[idx])
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_w = tc_pr.first_child_found_in("w:tcW")
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    tc_pr.append(tc_w)
                tc_w.set(qn("w:w"), str(int(widths_in[idx] * 1440)))
                tc_w.set(qn("w:type"), "dxa")


def add_table_from_df(document: Document, df: pd.DataFrame, columns: list[str], widths_in: list[float]) -> None:
    """Agrega una tabla Word desde un DataFrame pequeno."""
    table = document.add_table(rows=1, cols=len(columns))
    table.style = "Table Grid"
    header = table.rows[0].cells
    for idx, column in enumerate(columns):
        set_cell_text(header[idx], column, bold=True)
        set_cell_shading(header[idx], "E8EEF5")
        header[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for _, row in df[columns].iterrows():
        cells = table.add_row().cells
        for idx, column in enumerate(columns):
            value = row[column]
            if isinstance(value, float):
                if "rate" in column:
                    value = f"{value * 100:.2f}%"
                elif value.is_integer():
                    value = f"{int(value)}"
                else:
                    value = f"{value:.3f}"
            set_cell_text(cells[idx], value)
            if column in {"total_images", "images_with_LA", "accepted", "rejected"}:
                cells[idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    set_table_geometry(table, widths_in)
    document.add_paragraph()


def add_caption(document: Document, text: str) -> None:
    """Agrega caption sencillo para figura o tabla."""
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(85, 85, 85)


def add_figure(document: Document, image_path: Path, caption: str) -> None:
    """Agrega imagen si existe."""
    if not image_path.exists():
        document.add_paragraph(f"No se encontro la figura: {image_path.name}")
        return
    document.add_picture(str(image_path), width=Inches(5.9))
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(document, caption)


def configure_styles(document: Document) -> None:
    """Aplica preset standard_business_brief."""
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True


def pct(value: float) -> str:
    """Formato porcentaje."""
    return f"{value * 100:.2f}%"


def shorten_filename(value: str, max_len: int = 52) -> str:
    """Acorta nombres largos para tablas del Word."""
    value = str(value)
    if len(value) <= max_len:
        return value
    return value[: max_len - 3] + "..."


def main() -> None:
    """Construye el informe tecnico."""
    REPORTS_ROOT.mkdir(parents=True, exist_ok=True)

    metrics = pd.read_csv(METRICS_CSV)
    analysis = pd.read_csv(ANALYSIS_CSV)
    examples = pd.read_csv(EXAMPLES_CSV)
    visual_cases = pd.read_csv(VISUAL_CASES_CSV) if VISUAL_CASES_CSV.exists() else pd.DataFrame()
    thresholds = pd.read_csv(THRESHOLD_CSV)
    quality_group = pd.read_csv(QUALITY_GROUP_CSV)
    roboflow = pd.read_csv(ROBOFLOW_AUDIT_CSV)
    mask_report = pd.read_csv(MASK_REPORT_CSV)

    by_quality = analysis[analysis["section"] == "by_quality"].copy()
    by_split = analysis[analysis["section"] == "by_split"].copy()
    by_patient = analysis[analysis["section"] == "by_patient"].copy()

    total_images = len(metrics)
    images_with_la = int(pd.to_numeric(metrics["has_la"]).sum())
    accepted = int(pd.to_numeric(metrics["acceptable_lumen_threshold_rule"]).sum())
    rejected = total_images - accepted

    document = Document()
    configure_styles(document)

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("Reporte tecnico del pipeline longitudinal")
    run.font.name = "Calibri"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = document.add_paragraph()
    subtitle.add_run("Development of an Artificial Vision System for the Standardized Acquisition of Clinically Relevant Hepatic Ultrasound Images").italic = True
    document.add_paragraph("Fecha de generacion: 23 de junio de 2026")

    document.add_heading("1. Objetivo del analisis", level=1)
    document.add_paragraph(
        "Este documento resume el procesamiento computacional realizado para la vista longitudinal. "
        "El flujo parte del dataset COCO exportado desde Roboflow, genera mascaras binarias por clase, "
        "aplica dilatacion sobre el lumen anecoico, intersecta la mascara resultante con la ROI ecografica "
        "y calcula metricas de textura GLCM e intensidad para proponer una regla inicial de aceptabilidad."
    )

    document.add_heading("2. Estado del dataset y anotaciones", level=1)
    split_summary = (
        roboflow.groupby("split")
        .agg(image_count=("image_count", "sum"), annotations_ROI=("annotations_ROI", "sum"), annotations_Higado=("annotations_Higado", "sum"), annotations_LA=("annotations_LA", "sum"))
        .reset_index()
    )
    add_table_from_df(document, split_summary, ["split", "image_count", "annotations_ROI", "annotations_Higado", "annotations_LA"], [1.0, 1.2, 1.4, 1.5, 1.2])
    add_caption(document, "Tabla 1. Auditoria del dataset COCO exportado desde Roboflow.")

    document.add_paragraph(
        "El dataset longitudinal COCO contiene 909 imagenes de P001, P002 y P003. "
        "ROI e Higado aparecen en todas las imagenes, mientras que LA aparece solo cuando fue visible y anotable. "
        "Esto es metodologicamente importante porque evita inventar lumen en imagenes de baja calidad."
    )

    document.add_heading("3. Flujo computacional implementado", level=1)
    steps = [
        "Auditoria del COCO y confirmacion de clases ROI, Higado y LA.",
        "Conversion de poligonos COCO a mascaras binarias PNG por split y clase.",
        "Generacion de overlays de control visual para revisar correspondencia imagen-mascara.",
        "Dilatacion morfologica de LA y operacion AND con ROI.",
        "Extraccion de area, intensidad media, desviacion estandar y metricas GLCM.",
        "Definicion de umbrales iniciales usando imagenes clear de train y valid.",
    ]
    for step in steps:
        document.add_paragraph(step, style="List Bullet")

    document.add_heading("4. Umbrales iniciales de aceptabilidad", level=1)
    add_table_from_df(document, thresholds, ["parameter", "value", "method"], [2.0, 1.4, 3.1])
    add_caption(document, "Tabla 2. Umbrales iniciales calculados desde imagenes clear de train/valid.")
    document.add_paragraph(
        "La regla inicial acepta una imagen si existe LA, si el area del lumen supera el minimo calculado, "
        "y si la desviacion estandar y entropia quedan por debajo de los limites de referencia. "
        "Estos umbrales son un punto de partida cuantitativo, no una conclusion clinica final."
    )

    document.add_heading("5. Resultados generales de aceptabilidad", level=1)
    key_table = pd.DataFrame(
        [
            ["Imagenes analizadas", total_images],
            ["Imagenes con LA anotado", images_with_la],
            ["Imagenes aceptadas", accepted],
            ["Imagenes rechazadas", rejected],
            ["Tasa global de aceptacion", pct(accepted / total_images)],
        ],
        columns=["Indicador", "Valor"],
    )
    add_table_from_df(document, key_table, ["Indicador", "Valor"], [3.4, 2.4])
    add_caption(document, "Tabla 3. Resumen global del pipeline longitudinal.")

    document.add_heading("5.1 Aceptabilidad por calidad", level=2)
    add_table_from_df(document, by_quality, ["group", "total_images", "images_with_LA", "accepted", "rejected", "acceptance_rate", "acceptance_rate_with_LA"], [0.9, 0.9, 1.0, 0.8, 0.8, 1.05, 1.25])
    add_caption(document, "Tabla 4. Resultados por calidad manual.")
    add_figure(document, FIGURES_ROOT / "acceptance_rate_by_quality.png", "Figura 1. Tasa de aceptacion por calidad.")

    document.add_heading("5.2 Aceptabilidad por split", level=2)
    add_table_from_df(document, by_split, ["group", "total_images", "images_with_LA", "accepted", "rejected", "acceptance_rate", "acceptance_rate_with_LA"], [0.9, 0.9, 1.0, 0.8, 0.8, 1.05, 1.25])
    add_caption(document, "Tabla 5. Resultados por split interno de Roboflow.")
    add_figure(document, FIGURES_ROOT / "acceptance_rate_by_split.png", "Figura 2. Tasa de aceptacion por split.")

    document.add_heading("5.3 Aceptabilidad por paciente", level=2)
    add_table_from_df(document, by_patient, ["group", "total_images", "images_with_LA", "accepted", "rejected", "acceptance_rate", "acceptance_rate_with_LA"], [0.9, 0.9, 1.0, 0.8, 0.8, 1.05, 1.25])
    add_caption(document, "Tabla 6. Resultados por paciente dentro del dataset anotado.")
    add_figure(document, FIGURES_ROOT / "acceptance_rate_by_patient.png", "Figura 3. Tasa de aceptacion por paciente.")

    document.add_heading("6. Distribucion de metricas por calidad", level=1)
    add_figure(document, FIGURES_ROOT / "boxplot_std_by_quality.png", "Figura 4. Desviacion estandar del lumen por calidad.")
    add_figure(document, FIGURES_ROOT / "boxplot_entropy_by_quality.png", "Figura 5. Entropia GLCM por calidad.")
    add_figure(document, FIGURES_ROOT / "boxplot_contrast_by_quality.png", "Figura 6. Contraste GLCM por calidad.")

    document.add_heading("7. Ejemplos seleccionados para revision", level=1)
    examples_doc = examples.copy()
    examples_doc["filename_short"] = examples_doc["filename"].apply(shorten_filename)
    examples_doc["accepted"] = examples_doc["acceptable_lumen_threshold_rule"]
    examples_doc = examples_doc[["example_group", "split", "patient", "quality", "has_la", "accepted", "la_area_px", "la_std_intensity", "glcm_entropy", "filename_short"]]
    add_table_from_df(document, examples_doc, ["example_group", "split", "patient", "quality", "has_la", "accepted", "la_area_px", "la_std_intensity", "glcm_entropy", "filename_short"], [1.15, 0.55, 0.55, 0.65, 0.45, 0.55, 0.75, 0.75, 0.75, 1.25])
    add_caption(document, "Tabla 7. Ejemplos representativos; los nombres completos estan en longitudinal_acceptability_examples.csv.")

    document.add_heading("7.1 Casos visuales representativos revisados manualmente", level=2)
    document.add_paragraph(
        "Se incorporaron cuatro ejemplos revisados visualmente para separar dos conceptos: "
        "una imagen puede pertenecer nominalmente a la vista longitudinal, pero no necesariamente contener un lumen anecoico "
        "clinicamente confiable. Para esta tesis, se considera dudoso todo caso con circulo o region negra sin borde blanco "
        "claramente definido, porque no permite validar de forma defendible el lumen anecoico."
    )
    document.add_paragraph(
        "P001 representa una adquisicion longitudinal clara con LA visible. P002 y P003 representan vistas longitudinales "
        "nominales donde LA no se observa de manera confiable. El ejemplo blurry se mantiene como rechazado porque la definicion "
        "global y los bordes anatomicos no son suficientes para analisis posterior."
    )
    if not visual_cases.empty:
        visual_cases_doc = visual_cases.copy()
        visual_cases_doc["image_name"] = visual_cases_doc["image_path"].apply(lambda value: shorten_filename(Path(str(value)).name, 42))
        add_table_from_df(
            document,
            visual_cases_doc,
            ["case_id", "patient", "quality", "visual_category", "file_exists"],
            [1.35, 0.55, 0.65, 2.8, 0.65],
        )
        add_caption(document, "Tabla 8. Casos visuales definidos para interpretar aceptabilidad longitudinal.")
    add_figure(document, VISUAL_CASES_FIGURE, "Figura 7. Casos representativos de P001, P002 y P003 para sustentar la interpretacion visual.")

    document.add_heading("8. Interpretacion tecnica", level=1)
    document.add_paragraph(
        "La tasa global de aceptacion fue baja porque la regla exige que LA exista. "
        "Muchas imagenes no tienen LA anotado, especialmente cuando el lumen no fue visible o no era defendible anotarlo."
    )
    document.add_paragraph(
        "El comportamiento por paciente muestra una concentracion fuerte de LA anotado en P001. "
        "Tras la revision visual de overlays, esta concentracion se interpreta principalmente como una diferencia real "
        "en la calidad de adquisicion entre pacientes: las vistas longitudinales de P002 y P003 no quedaron adquiridas "
        "con la misma claridad anatomica que P001."
    )
    document.add_heading("8.1 Interpretacion clinica y metodologica del desbalance de LA", level=2)
    document.add_paragraph(
        "Este hallazgo no invalida el pipeline. Al contrario, respalda la motivacion central de la tesis: "
        "la adquisicion ecografica hepatica depende de manera importante del operador, de la orientacion del transductor, "
        "de la ventana acustica y de la capacidad de mantener una vista anatomica clinicamente util."
    )
    document.add_paragraph(
        "Aunque P001, P002 y P003 pertenecen nominalmente a la vista longitudinal, no todas las imagenes longitudinales "
        "contienen el mismo nivel de informacion clinica. En P001 el lumen anecoico fue visible y anotable con mayor frecuencia, "
        "mientras que en P002 y P003 el lumen casi no aparece de forma suficientemente clara. Por tanto, la etiqueta de vista "
        "por si sola no garantiza que la imagen sea informativa para evaluar la vena porta o el lumen esperado."
    )
    document.add_paragraph(
        "Esta observacion justifica el uso de metricas objetivas y reglas de aceptabilidad. El sistema no solo debe reconocer "
        "que una imagen pertenece a una vista longitudinal, sino tambien evaluar si contiene estructuras clinicamente relevantes "
        "con calidad suficiente para analisis posterior."
    )
    document.add_paragraph(
        "La regla inicial es util como criterio cuantitativo preliminar, pero debe validarse visualmente con overlays y casos dudosos. "
        "En particular, los casos blurry aceptados deben revisarse para confirmar si corresponden a lumen real o a una anotacion que podria ser demasiado permisiva. "
        "Los casos con una region oscura circular sin pared hiperecogenica definida deben tratarse como dudosos, aun si la region parece anecoica."
    )

    document.add_heading("9. Archivos generados", level=1)
    files = [
        "outputs/reports/longitudinal_acceptability_analysis.csv",
        "outputs/reports/longitudinal_acceptability_summary.md",
        "outputs/reports/longitudinal_acceptability_examples.csv",
        "outputs/reports/longitudinal_visual_case_examples.csv",
        "outputs/metrics/glcm_longitudinal_metrics_with_thresholds.csv",
        "outputs/reports/threshold_summary.csv",
        "outputs/reports/quality_group_summary.csv",
        "outputs/figures/acceptance_rate_by_quality.png",
        "outputs/figures/acceptance_rate_by_patient.png",
        "outputs/figures/acceptance_rate_by_split.png",
        "outputs/figures/casos_representativos_longitudinal.png",
    ]
    for item in files:
        document.add_paragraph(item, style="List Bullet")

    document.add_heading("10. Siguiente paso recomendado", level=1)
    document.add_paragraph(
        "Revisar visualmente los overlays y los 20 ejemplos seleccionados. "
        "Con la revision realizada, el siguiente paso es consolidar una version final de la regla de aceptabilidad "
        "incluyendo un criterio visual adicional para casos dudosos: region negra sin pared blanca definida no debe contar como LA valida. "
        "Luego se puede preparar la comparacion local de U-Net, DeepLabV3+ y SegFormer priorizando la clase LA."
    )

    document.save(DOCX_PATH)
    print(f"Documento generado: {DOCX_PATH}")


if __name__ == "__main__":
    main()


