"""Genera documento tecnico Word/Markdown para segmentacion longitudinal."""
from __future__ import annotations
from datetime import date
import os
from pathlib import Path
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

TESIS_ROOT = Path(os.environ.get("THESIS_PROJECT_ROOT", Path(__file__).resolve().parents[1]))
OUT = TESIS_ROOT / "outputs" / "segmentation_training"
REPORTS = OUT / "reports"
FIGURES = OUT / "figures"
DOCX = REPORTS / "documento_tecnico_segmentacion_longitudinal.docx"
MD = REPORTS / "documento_tecnico_segmentacion_longitudinal.md"

ARCH_REFS = [
    ["U-Net", "Ronneberger, Fischer y Brox, 2015", "Baseline clasico de imagen medica; encoder-decoder con buena localizacion espacial.", "https://arxiv.org/abs/1505.04597"],
    ["DeepLabV3+", "Chen et al., 2018", "CNN moderna con contexto multiescala y decoder para refinar bordes.", "https://arxiv.org/abs/1802.02611"],
    ["SegFormer", "Xie et al., 2021", "Transformer eficiente con encoder jerarquico y decoder MLP ligero.", "https://arxiv.org/abs/2105.15203"],
]


def fmt(v, digits=4):
    if v is None or pd.isna(v):
        return "NA"
    try:
        return f"{float(v):.{digits}f}"
    except Exception:
        return str(v)


def read_csv(name: str) -> pd.DataFrame:
    path = REPORTS / name
    return pd.read_csv(path) if path.exists() else pd.DataFrame()


def shade(cell, fill="F2F4F7"):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell(cell_obj, text, bold=False, size=7.5):
    cell_obj.text = ""
    p = cell_obj.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(str(text)) < 18 else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(size)
    cell_obj.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def table(doc, title, headers, rows):
    doc.add_paragraph(title, style="Heading 3")
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.style = "Table Grid"
    for i, h in enumerate(headers):
        cell(t.rows[0].cells[i], h, True, 8)
        shade(t.rows[0].cells[i])
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cell(cells[i], value, False, 7.2)
    doc.add_paragraph("")


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def figure(doc, path: Path, caption: str, width=5.8):
    if not path.exists():
        doc.add_paragraph(f"Figura no encontrada: {path}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    c = doc.add_paragraph(caption)
    c.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in c.runs:
        r.font.size = Pt(9)
        r.font.italic = True


def configure(doc):
    sec = doc.sections[0]
    sec.top_margin = sec.bottom_margin = sec.left_margin = sec.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    for name, size, color in [("Heading 1", 16, RGBColor(46,116,181)), ("Heading 2", 13, RGBColor(46,116,181)), ("Heading 3", 12, RGBColor(31,77,120))]:
        s = doc.styles[name]
        s.font.name = "Calibri"
        s.font.size = Pt(size)
        s.font.color.rgb = color
        s.font.bold = True


def audit_rows(df):
    """Adapta la auditoria real de COCO separados a una tabla compacta."""
    if df.empty:
        return []
    rows = []
    for _, r in df.iterrows():
        rows.append([
            r.get("dataset_class_expected", ""),
            r.get("split", ""),
            int(r.get("image_count_coco", 0)),
            int(r.get("annotation_count_target", 0)),
            r.get("detected_target_category", r.get("category_names", "")),
        ])
    return rows


def result_rows(df):
    rows = []
    for _, r in df.iterrows():
        rows.append([
            r.get("class_name", ""),
            r.get("architecture", ""),
            fmt(r.get("test_dice")),
            fmt(r.get("test_iou")),
            fmt(r.get("test_precision")),
            fmt(r.get("test_recall")),
            fmt(r.get("test_positive_dice")),
            fmt(r.get("empty_gt_false_positive_rate")),
            fmt(r.get("inference_time_s_per_frame"), 5),
            f"{int(r.get('parameter_count', 0)):,}" if not pd.isna(r.get("parameter_count", None)) else "NA",
        ])
    return rows


def best_rows(df):
    rows = []
    for _, r in df.iterrows():
        rows.append([
            r.get("class_name", ""),
            r.get("architecture", ""),
            fmt(r.get("test_dice")),
            fmt(r.get("test_iou")),
            fmt(r.get("test_positive_dice")),
            fmt(r.get("empty_gt_false_positive_rate")),
            Path(str(r.get("checkpoint_path", ""))).name,
        ])
    return rows


def df_markdown(df: pd.DataFrame) -> str:
    """Convierte DataFrame a Markdown sin depender de tabulate."""
    if df.empty:
        return "Sin datos."
    headers = [str(c) for c in df.columns]
    lines = ["| " + " | ".join(headers) + " |", "| " + " | ".join(["---"] * len(headers)) + " |"]
    for _, row in df.iterrows():
        values = [str(row.get(c, "")) for c in df.columns]
        lines.append("| " + " | ".join(v.replace("\n", " ") for v in values) + " |")
    return "\n".join(lines)

def write_markdown(results, best, audit):
    lines = [
        "# Documento tecnico de segmentacion longitudinal",
        "",
        f"Fecha: {date.today().isoformat()}",
        "",
        "## Origen de las arquitecturas",
    ]
    for arch, origin, use, url in ARCH_REFS:
        lines.append(f"- **{arch}**: {origin}. {use} Referencia: {url}")
    lines += ["", "## Auditoria COCO", df_markdown(audit) if not audit.empty else "Sin auditoria CSV."]
    lines += ["", "## Ranking de resultados", df_markdown(results) if not results.empty else "Sin ranking CSV."]
    lines += ["", "## Modelos seleccionados", df_markdown(best) if not best.empty else "Sin mejores modelos CSV."]
    lines += [
        "",
        "## Fallos y correcciones",
        "- DeepLabV3+ para LA selecciono inicialmente una solucion casi vacia por usar Dice global.",
        "- Se agregaron metricas separadas para imagenes positivas y para imagenes vacias.",
        "- Para LA se uso combined_la_score = positive_dice - empty_false_positive_rate.",
        "- El evaluador reporta falsos positivos sobre imagenes sin LA y metricas positivas en imagenes con LA.",
        "- El pipeline de inferencia calcula GLCM real sobre la mascara LA y guarda CSV por frame.",
    ]
    MD.write_text("\n".join(lines), encoding="utf-8")


def build():
    REPORTS.mkdir(parents=True, exist_ok=True)
    results = read_csv("architecture_ranking_by_class.csv")
    best = read_csv("best_models_by_class.csv")
    audit = read_csv("coco_separated_audit.csv")
    write_markdown(results, best, audit)

    doc = Document()
    configure(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Reporte tecnico de segmentacion longitudinal")
    r.bold = True
    r.font.size = Pt(20)
    r.font.color.rgb = RGBColor(11, 37, 69)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Sistema de vision artificial para adquisicion estandarizada de ecografias hepaticas").italic = True
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Fecha de generacion: {date.today().isoformat()}")

    doc.add_heading("1. Objetivo de esta fase", level=1)
    doc.add_paragraph("Esta fase evalua tres arquitecturas de segmentacion para la vista longitudinal hepatica. La meta es obtener modelos locales reproducibles para detectar ROI, Higado y LA, y preparar una inferencia frame por frame para una GUI futura.")

    doc.add_heading("2. De donde salen las tres arquitecturas", level=1)
    doc.add_paragraph("Las arquitecturas fueron seleccionadas desde literatura reconocida de segmentacion semantica: una referencia clasica de imagen medica, una CNN multiescala moderna y una alternativa Transformer eficiente.")
    table(doc, "Tabla 1. Referencias de las arquitecturas.", ["Arquitectura", "Origen", "Uso en tesis", "Referencia"], ARCH_REFS)

    doc.add_heading("3. Dataset y auditoria", level=1)
    doc.add_paragraph("Se usaron datasets COCO separados descargados desde Roboflow: ROI_COCO, Higado_COCO y LA_COCO. Roboflow se conserva como herramienta de anotacion/exportacion, no como dependencia final de inferencia.")
    table(doc, "Tabla 2. Auditoria de datasets COCO separados.", ["Dataset", "Split", "Imagenes", "Anotaciones", "Clase"], audit_rows(audit))

    doc.add_heading("4. Flujo realizado", level=1)
    bullets(doc, [
        "Auditoria de ROI_COCO, Higado_COCO y LA_COCO en train/valid/test.",
        "Implementacion de Dataset PyTorch para COCO segmentation y conversion de poligonos a mascaras binarias.",
        "Resize local a 512x512 dentro del entrenamiento, sin modificar los datasets Roboflow limpios.",
        "Entrenamiento de 9 combinaciones: U-Net, DeepLabV3+ y SegFormer para ROI, Higado y LA.",
        "Evaluacion en test con Dice, IoU, precision, recall, F1, tiempo por frame y parametros.",
        "Generacion de checkpoints, logs, curvas, overlays, CSV, Markdown y modelos finales.",
    ])

    doc.add_heading("5. Fallos encontrados y correcciones", level=1)
    doc.add_paragraph("El principal problema aparecio en la clase LA por desbalance: muchas imagenes no tienen lumen anotado. En ese escenario, una metrica global puede favorecer modelos que predicen mascaras vacias.")
    bullets(doc, [
        "DeepLabV3+ LA selecciono inicialmente una solucion casi vacia al optimizar Dice global.",
        "Se respaldaron los pesos iniciales y se reentreno LA con una metrica mas adecuada.",
        "Se agrego combined_la_score = positive_dice - empty_false_positive_rate para seleccionar checkpoints de LA.",
        "El evaluador ahora separa imagenes positivas, imagenes vacias y falsos positivos sobre vacias.",
        "El comparador final usa criterio especial para LA y criterio global para ROI/Higado.",
        "El pipeline de inferencia reemplazo la entropia placeholder por GLCM real sobre la mascara LA.",
    ])

    doc.add_heading("6. Resultados completos", level=1)
    table(doc, "Tabla 3. Resultados de test por arquitectura y clase.", ["Clase", "Arquitectura", "Dice", "IoU", "Precision", "Recall", "Dice positivo", "FP vacias", "s/frame", "Parametros"], result_rows(results))
    table(doc, "Tabla 4. Modelos seleccionados para inferencia final.", ["Clase", "Modelo", "Dice", "IoU", "Dice positivo", "FP vacias", "Checkpoint"], best_rows(best))

    doc.add_heading("7. Graficas", level=1)
    for filename, caption in [
        ("comparison_test_dice.png", "Figura 1. Dice global por arquitectura y clase."),
        ("comparison_test_iou.png", "Figura 2. IoU global por arquitectura y clase."),
        ("comparison_test_positive_dice.png", "Figura 3. Dice positivo, especialmente relevante para LA."),
        ("comparison_inference_time_s_per_frame.png", "Figura 4. Tiempo promedio de inferencia por frame."),
    ]:
        figure(doc, FIGURES / filename, caption)

    doc.add_heading("8. Curvas de entrenamiento", level=1)
    doc.add_paragraph("Las curvas permiten verificar convergencia y estabilidad por clase y arquitectura.")
    for class_name in ["ROI", "Higado", "LA"]:
        doc.add_heading(f"Curvas para {class_name}", level=2)
        for arch in ["unet", "deeplabv3", "segformer"]:
            figure(doc, FIGURES / class_name / f"{arch}_{class_name.lower()}_curves.png", f"Curva de entrenamiento: {arch} - {class_name}.", width=5.5)

    doc.add_heading("9. Pipeline de inferencia preparado", level=1)
    doc.add_paragraph("El script 08_inferencia_video_longitudinal_base.py procesa video frame por frame, predice ROI/Higado/LA, calcula areas, proporcion higado/ROI, desviacion estandar, GLCM y decision textual para usuario no experto.")
    table(doc, "Tabla 5. Mensajes propuestos para GUI.", ["Condicion computacional", "Mensaje para usuario"], [
        ["ROI ausente o imagen mayormente negra", "Ninguna estructura visible. Revisar contacto y gel."],
        ["ROI visible pero Higado pequeno o ausente", "Higado parcialmente visible. Mover la sonda para centrarlo."],
        ["Higado visible pero LA ausente o textura fuera de rango", "Referencia anatomica insuficiente. Ajustar inclinacion o posicion."],
        ["ROI, Higado y LA visibles con area/textura aceptables", "Higado visible. Mantener posicion y capturar imagen."],
    ])
    doc.add_paragraph("Comando base desde Visual Studio:")
    cmd = doc.add_paragraph()
    run = cmd.add_run('cd "<PROJECT_ROOT>"\npython ".\\Codigos_Entrenamiento_Segmentacion\\08_inferencia_video_longitudinal_base.py" --video_path "ruta\\al\\video.mp4" --frame_stride 5 --save_overlays')
    run.font.name = "Consolas"
    run.font.size = Pt(9)

    doc.add_heading("10. Resumen de lo mas importante", level=1)
    bullets(doc, [
        "ROI quedo mejor con DeepLabV3+ por una diferencia minima de Dice/IoU.",
        "Higado quedo mejor con U-Net por mejor Dice/IoU y buena velocidad.",
        "LA quedo mejor con U-Net por mejor equilibrio entre Dice positivo y menor tasa de falsos positivos.",
        "Las metricas globales para LA deben interpretarse con cautela por el desbalance de imagenes positivas.",
        "Ya existe una base local reproducible para inferencia longitudinal sobre video.",
    ])

    doc.add_heading("11. Referencias", level=1)
    refs = [
        "Ronneberger O., Fischer P., Brox T. U-Net: Convolutional Networks for Biomedical Image Segmentation. MICCAI 2015. https://arxiv.org/abs/1505.04597",
        "Chen L.-C., Zhu Y., Papandreou G., Schroff F., Adam H. Encoder-Decoder with Atrous Separable Convolution for Semantic Image Segmentation. ECCV 2018. https://arxiv.org/abs/1802.02611",
        "Xie E., Wang W., Yu Z., Anandkumar A., Alvarez J. M., Luo P. SegFormer: Simple and Efficient Design for Semantic Segmentation with Transformers. NeurIPS 2021. https://arxiv.org/abs/2105.15203",
    ]
    for ref in refs:
        doc.add_paragraph(ref, style="List Number")

    doc.save(DOCX)
    print(f"Documento Word generado: {DOCX}")
    print(f"Documento Markdown generado: {MD}")


if __name__ == "__main__":
    build()


